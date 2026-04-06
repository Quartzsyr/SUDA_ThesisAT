from __future__ import annotations
import re
from dataclasses import dataclass
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百千万零两\d]+章(?:\s+|　+|$)")
SECTION_RE = re.compile(r"^第\d+\.\d+节(?:\s+|　+|$)")
SUBSECTION_RE = re.compile(r"^第\d+\.\d+\.\d+节(?:\s+|　+|$)")
NUMERIC_RE = re.compile(r"^(\d+(?:\.\d+){0,2})\s+")
CAPTION_RE = re.compile(r"^(?:图|Figure|表|Table)\s*\d+(?:[-.－]\d+)?\s*", re.IGNORECASE)
FORMULA_PLACEHOLDER_RE = re.compile(r"^\[公式\s*(\d+)\]$")
STYLE_LEVELS = {
    "一级标题": 1,
    "二级标题": 2,
    "三级标题": 3,
}
STYLE_ID_LEVELS = {"ae": 1, "af0": 2, "af2": 3}
SPECIAL_H1 = {"摘要", "abstract", "目录", "参考文献", "致谢", "结论", "附录"}
@dataclass
class ImageBinding:
    caption_key: str
    image_bytes: bytes
    width_emu: int | None = None
    height_emu: int | None = None
def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())
def _paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath('.//*[local-name()="drawing"]'))
def _paragraph_has_math(paragraph) -> bool:
    return bool(paragraph._element.xpath('.//*[local-name()="oMath"]')) or bool(paragraph._element.xpath('.//*[local-name()="oMathPara"]'))
def _is_caption_text(text: str) -> bool:
    return bool(CAPTION_RE.match(_normalize(text)))
def caption_key(text: str) -> str:
    normalized = _normalize(text)
    normalized = CAPTION_RE.sub("", normalized, count=1)
    return normalized.strip().lower()
def detect_heading_level(style_name: str, style_id: str, text: str) -> int | None:
    style_name = (style_name or "").strip()
    style_id = (style_id or "").strip()
    text = _normalize(text)
    lowered_style = style_name.lower()
    lowered_text = text.lower()
    if not text:
        return None
    if style_name in STYLE_LEVELS:
        return STYLE_LEVELS[style_name]
    if style_id in STYLE_ID_LEVELS:
        return STYLE_ID_LEVELS[style_id]
    if "heading 1" in lowered_style or "标题 1" in style_name:
        return 1
    if "heading 2" in lowered_style or "标题 2" in style_name:
        return 2
    if "heading 3" in lowered_style or "标题 3" in style_name:
        return 3
    if lowered_text in SPECIAL_H1:
        return 1
    if CHAPTER_RE.match(text):
        return 1
    if SUBSECTION_RE.match(text):
        return 3
    if SECTION_RE.match(text):
        return 2
    match = NUMERIC_RE.match(text)
    if match:
        return min(match.group(1).count(".") + 1, 3)
    return None
def _is_cover_or_statement(text: str) -> bool:
    normalized = _normalize(text)
    keywords = [
        "本科毕业设计（论文）独创性声明",
        "本科毕业设计（论文）使用授权声明",
        "作者签名：",
        "论文作者签名：",
        "导师签名：",
        "日 期：",
        "日  期：",
        "涉密设计（论文）",
        "非涉密设计（论文）",
    ]
    return any(token in normalized for token in keywords)
def _is_title_candidate(text: str) -> bool:
    text = _normalize(text)
    if len(text) < 6:
        return False
    if text.lower() in SPECIAL_H1:
        return False
    if CHAPTER_RE.match(text) or SECTION_RE.match(text) or SUBSECTION_RE.match(text):
        return False
    if "声明" in text or "签名" in text:
        return False
    return True
def _extract_image_payloads(paragraph) -> list[ImageBinding]:
    bindings: list[ImageBinding] = []
    for container in paragraph._element.xpath('.//*[local-name()="inline" or local-name()="anchor"]'):
        blips = container.xpath('.//*[local-name()="blip"]')
        if not blips:
            continue
        embed = blips[0].get(qn('r:embed'))
        if not embed:
            continue
        related = paragraph.part.related_parts.get(embed)
        if related is None:
            continue
        width_emu = None
        height_emu = None
        extents = container.xpath('.//*[local-name()="extent"]')
        if extents:
            width_emu = int(extents[0].get('cx')) if extents[0].get('cx') else None
            height_emu = int(extents[0].get('cy')) if extents[0].get('cy') else None
        bindings.append(ImageBinding(caption_key="", image_bytes=related.blob, width_emu=width_emu, height_emu=height_emu))
    return bindings
def extract_captioned_images(path: str | Path) -> dict[str, list[ImageBinding]]:
    doc = Document(str(path))
    mapping: dict[str, list[ImageBinding]] = {}
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        payloads = _extract_image_payloads(paragraph)
        if not payloads:
            continue
        caption_text = _normalize(paragraph.text)
        if not _is_caption_text(caption_text):
            for look_ahead in range(index + 1, min(index + 4, len(paragraphs))):
                candidate = paragraphs[look_ahead]
                candidate_text = _normalize(candidate.text)
                if not candidate_text:
                    continue
                if _paragraph_has_drawing(candidate):
                    break
                if _is_caption_text(candidate_text):
                    caption_text = candidate_text
                    break
                break
        if not _is_caption_text(caption_text):
            continue
        key = caption_key(caption_text)
        queue = mapping.setdefault(key, [])
        for payload in payloads:
            payload.caption_key = key
            queue.append(payload)
    return mapping
def extract_formula_paragraphs(path: str | Path) -> dict[str, object]:
    doc = Document(str(path))
    formulas: dict[str, object] = {}
    index = 0
    for paragraph in doc.paragraphs:
        if _paragraph_has_math(paragraph) and not _normalize(paragraph.text):
            index += 1
            formulas[f"[公式 {index}]"] = paragraph._element
    return formulas
def document_to_markup(path: str | Path) -> tuple[str, str]:
    doc = Document(str(path))
    suggested_title = Path(path).stem
    collected: list[str] = []
    first_heading_seen = False
    pending_title: str | None = None
    formula_index = 0
    for para in doc.paragraphs:
        text = _normalize(para.text)
        has_math = _paragraph_has_math(para)
        if not text and not has_math:
            continue
        style = getattr(para, "style", None)
        style_name = getattr(style, "name", "")
        style_id = getattr(style, "style_id", "")
        level = detect_heading_level(style_name, style_id, text)
        if not first_heading_seen:
            if level == 1:
                first_heading_seen = True
                if pending_title:
                    suggested_title = pending_title
                collected.append(f"# {text}")
                continue
            if _is_cover_or_statement(text):
                continue
            if _is_title_candidate(text):
                pending_title = text
            continue
        if has_math and not text:
            formula_index += 1
            collected.append(f"[\u516c\u5f0f {formula_index}]")
            continue
        if level:
            collected.append(f"{'#' * level} {text}")
        else:
            collected.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                collected.append(" | ".join(cells))
    return "\n\n".join(collected), suggested_title
