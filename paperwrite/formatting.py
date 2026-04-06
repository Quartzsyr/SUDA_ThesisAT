from __future__ import annotations

import html
import re
from copy import deepcopy
from io import BytesIO
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

from paperwrite.docx_loader import caption_key, extract_captioned_images, extract_formula_paragraphs


FONT_DIR = Path(r"C:\Windows\Fonts")
registerFont(TTFont("SimSun", str(FONT_DIR / "simsun.ttc")))
registerFont(TTFont("SimHei", str(FONT_DIR / "simhei.ttf")))
registerFont(TTFont("TimesNewRoman", str(FONT_DIR / "times.ttf")))
registerFont(TTFont("TimesNewRoman-Bold", str(FONT_DIR / "timesbd.ttf")))

CHAPTER_HEADING_RE = re.compile(r"^第[一二三四五六七八九十百千万零两\d]+章")
NUMERIC_CHAPTER_RE = re.compile(r"^(\d+)\s+")
FIGURE_CAPTION_RE = re.compile(r"^(图|Figure)\s*\d+(?:[-.－]\d+)?\s*", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(r"^(表|Table)\s*\d+(?:[-.－]\d+)?\s*", re.IGNORECASE)
REFERENCE_PREFIX_RE = re.compile(r"^\s*(?:\[\d+\]|\d+[.、])\s*")
SPECIAL_HEADINGS = {"摘要", "abstract", "目录", "参考文献", "致谢", "结论", "附录"}
ASSET_DIR = Path(__file__).resolve().parent.parent / "assets"
UNIVERSITY_IMAGE = ASSET_DIR / "suda-cover-header.png"


@dataclass
class FormatOptions:
    school_name: str = "\u82cf\u5dde\u5927\u5b66"
    header_text: str = "\u82cf\u5dde\u5927\u5b66\u672c\u79d1\u751f\u6bd5\u4e1a\u8bbe\u8ba1\uff08\u8bba\u6587\uff09"
    thesis_title: str = ""
    college: str = ""
    year_grade: str = ""
    major: str = ""
    class_name: str = ""
    student_id: str = ""
    author_name: str = ""
    supervisor: str = ""
    supervisor_title: str = ""
    date_text: str = ""
    include_cover: bool = True
    source_docx_path: str = ""


class ThesisDocTemplate(BaseDocTemplate):
    def __init__(self, *args, **kwargs):
        self._heading_toc_enabled = False
        super().__init__(*args, **kwargs)

    def afterFlowable(self, flowable):
        level = getattr(flowable, "toc_level", None)
        if level is None or not self._heading_toc_enabled:
            return
        text = getattr(flowable, "toc_text", "")
        page_num = self.page
        self.notify("TOCEntry", (level, text, page_num))


def _styles() -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    return {
        "body_cn": ParagraphStyle(
            "ThesisBodyCN",
            parent=styles["Normal"],
            fontName="SimSun",
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            firstLineIndent=24,
            spaceBefore=0,
            spaceAfter=0,
        ),
        "body_cn_no_indent": ParagraphStyle(
            "ThesisBodyCNNoIndent",
            parent=styles["Normal"],
            fontName="SimSun",
            fontSize=12,
            leading=18,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceBefore=0,
            spaceAfter=0,
        ),
        "body_en": ParagraphStyle(
            "ThesisBodyEN",
            parent=styles["Normal"],
            fontName="TimesNewRoman",
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            firstLineIndent=24,
            spaceBefore=0,
            spaceAfter=0,
        ),
        "body_en_no_indent": ParagraphStyle(
            "ThesisBodyENNoIndent",
            parent=styles["Normal"],
            fontName="TimesNewRoman",
            fontSize=12,
            leading=18,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceBefore=0,
            spaceAfter=0,
        ),
        "center_plain_en": ParagraphStyle(
            "CenterPlainEN",
            parent=styles["Normal"],
            fontName="TimesNewRoman",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "h1_cn": ParagraphStyle(
            "ThesisH1CN",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=15,
            leading=22,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceBefore=15.6,
            spaceAfter=15.6,
        ),
        "h1_en": ParagraphStyle(
            "ThesisH1EN",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=15,
            leading=22,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceBefore=15.6,
            spaceAfter=15.6,
        ),
        "h2_cn": ParagraphStyle(
            "ThesisH2CN",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=14,
            leading=20,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceBefore=15.6,
            spaceAfter=15.6,
        ),
        "h3_cn": ParagraphStyle(
            "ThesisH3CN",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=12,
            leading=18,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceBefore=15.6,
            spaceAfter=15.6,
        ),
        "toc_title": ParagraphStyle(
            "TOCTitle",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=15,
            leading=22,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceBefore=15.6,
            spaceAfter=15.6,
        ),
        "toc1": ParagraphStyle(
            "TOC1Style",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=14,
            leading=20,
            leftIndent=0,
            firstLineIndent=0,
        ),
        "toc2": ParagraphStyle(
            "TOC2Style",
            parent=styles["Normal"],
            fontName="SimSun",
            fontSize=14,
            leading=20,
            leftIndent=18,
            firstLineIndent=0,
        ),
        "toc3": ParagraphStyle(
            "TOC3Style",
            parent=styles["Normal"],
            fontName="SimSun",
            fontSize=14,
            leading=20,
            leftIndent=36,
            firstLineIndent=0,
        ),
        "coverLabel": ParagraphStyle(
            "CoverLabel",
            parent=styles["Normal"],
            fontName="SimSun",
            fontSize=14,
            leading=20,
            alignment=TA_LEFT,
            firstLineIndent=0,
        ),
        "coverValueCenter": ParagraphStyle(
            "CoverValueCenter",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            firstLineIndent=0,
        ),
        "coverValueLeft": ParagraphStyle(
            "CoverValueLeft",
            parent=styles["Normal"],
            fontName="SimHei",
            fontSize=14,
            leading=18,
            alignment=TA_LEFT,
            firstLineIndent=0,
        ),
    }


def _contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def _is_mostly_english(text: str) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    ascii_letters = sum(ch.isascii() for ch in letters)
    return ascii_letters / len(letters) > 0.8


def _to_roman(num: int) -> str:
    values = [(1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'), (100, 'c'), (90, 'xc'), (50, 'l'), (40, 'xl'), (10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i')]
    result = []
    for value, numeral in values:
        while num >= value:
            result.append(numeral)
            num -= value
    return ''.join(result) or 'i'


def _front_page(canvas, doc, header_text: str, cover_pages: int) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setFont("SimSun", 9)
    canvas.drawCentredString(width / 2, height - 2.1 * cm, header_text)
    canvas.setLineWidth(0.5)
    canvas.line(2.5 * cm, height - 2.35 * cm, width - 2.5 * cm, height - 2.35 * cm)
    logical_page = max(canvas.getPageNumber() - cover_pages, 1)
    canvas.drawCentredString(width / 2, 1.25 * cm, _to_roman(logical_page))
    canvas.restoreState()


def _body_page(canvas, doc, header_text: str, cover_pages: int, front_pages: int) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setFont("SimSun", 9)
    canvas.drawCentredString(width / 2, height - 2.1 * cm, header_text)
    canvas.setLineWidth(0.5)
    canvas.line(2.5 * cm, height - 2.35 * cm, width - 2.5 * cm, height - 2.35 * cm)
    logical_page = canvas.getPageNumber() - cover_pages - front_pages
    canvas.drawCentredString(width / 2, 1.25 * cm, f"-{logical_page}-")
    canvas.restoreState()


def _cover_story(options: FormatOptions, styles: dict[str, ParagraphStyle]) -> list:
    c_label = styles["coverLabel"]
    c_value_center = styles["coverValueCenter"]
    c_value_left = styles["coverValueLeft"]
    table_data = [
        [Paragraph("\u5b66\u9662(\u90e8)", c_label), Paragraph(html.escape(options.college or " "), c_value_center), "", "", ""],
        [Paragraph("\u9898   \u76ee", c_label), Paragraph(html.escape(options.thesis_title or " "), c_value_center), "", "", ""],
        ["", "", "", "", ""],
        [Paragraph("\u5e74  \u7ea7", c_label), Paragraph(html.escape(options.year_grade or " "), c_value_center), "", Paragraph("\u4e13\u4e1a", c_label), Paragraph(html.escape(options.major or " "), c_value_center)],
        [Paragraph("\u73ed  \u7ea7", c_label), Paragraph(html.escape(options.class_name or " "), c_value_center), "", Paragraph("\u5b66\u53f7", c_label), Paragraph(html.escape(options.student_id or " "), c_value_center)],
        [Paragraph("\u59d3  \u540d", c_label), Paragraph(html.escape(options.author_name or " "), c_value_center), "", "", ""],
        [Paragraph("\u6307\u5bfc\u8001\u5e08", c_label), Paragraph(html.escape(options.supervisor or " "), c_value_left), "", Paragraph("\u804c\u79f0", c_label), Paragraph(html.escape(options.supervisor_title or " "), c_value_center)],
        [Paragraph("\u8bba\u6587\u63d0\u4ea4\u65e5\u671f", c_label), Paragraph(html.escape(options.date_text or " "), c_value_center), "", "", ""],
    ]

    table = Table(table_data, colWidths=[2.8 * cm, 3.0 * cm, 1.8 * cm, 2.6 * cm, 4.8 * cm], hAlign="CENTER")
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "SimSun"),
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("SPAN", (1, 0), (4, 0)),
        ("SPAN", (1, 1), (4, 1)),
        ("SPAN", (0, 2), (4, 2)),
        ("SPAN", (1, 3), (2, 3)),
        ("SPAN", (1, 4), (2, 4)),
        ("SPAN", (1, 5), (4, 5)),
        ("SPAN", (1, 6), (2, 6)),
        ("SPAN", (1, 7), (4, 7)),
        ("LINEBELOW", (1, 0), (4, 1), 0.6, colors.black),
        ("LINEBELOW", (1, 3), (2, 4), 0.6, colors.black),
        ("LINEBELOW", (4, 3), (4, 4), 0.6, colors.black),
        ("LINEBELOW", (1, 5), (4, 5), 0.6, colors.black),
        ("LINEBELOW", (1, 6), (2, 6), 0.6, colors.black),
        ("LINEBELOW", (4, 6), (4, 6), 0.6, colors.black),
        ("LINEBELOW", (1, 7), (4, 7), 0.6, colors.black),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 2), (-1, 2), 16),
        ("BOTTOMPADDING", (0, 2), (-1, 2), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("ALIGN", (0, 0), (0, -1), "LEFT"),
        ("ALIGN", (3, 3), (3, 6), "LEFT"),
        ("ALIGN", (1, 6), (2, 6), "LEFT"),
    ]))

    story = [Spacer(1, 2.1 * cm)]
    if options.school_name == "苏州大学" and UNIVERSITY_IMAGE.exists():
        story.append(Image(str(UNIVERSITY_IMAGE), width=7.8 * cm, height=5.1 * cm, hAlign="CENTER"))
        story.append(Spacer(1, 1.1 * cm))
    else:
        school_style = ParagraphStyle(
            "CoverSchoolName",
            parent=styles["coverValueCenter"],
            fontName="SimHei" if _contains_cjk(options.school_name) else "TimesNewRoman-Bold",
            fontSize=24,
            leading=30,
            alignment=TA_CENTER,
        )
        title_style = ParagraphStyle(
            "CoverSchoolTitle",
            parent=styles["coverValueCenter"],
            fontSize=22,
            leading=28,
            alignment=TA_CENTER,
        )
        title_text = options.header_text.replace(options.school_name, "").strip("（）() ") or "本科毕业设计（论文）"
        story.append(Paragraph(html.escape(options.school_name), school_style))
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph(html.escape(title_text), title_style))
        story.append(Spacer(1, 1.0 * cm))
    story.append(table)
    return story


def _is_keyword_line(text: str) -> bool:
    lowered = text.strip().lower()
    return lowered.startswith("\u5173\u952e\u8bcd") or lowered.startswith("keywords")


def _looks_like_english_title(block: str) -> bool:
    if len(block) < 20:
        return False
    ascii_count = sum(1 for ch in block if ch.isascii() and (ch.isalpha() or ch.isspace()))
    return ascii_count >= len(block) * 0.65


def _split_front_and_body(markup_text: str) -> tuple[list[str], list[str]]:
    front: list[str] = []
    body: list[str] = []
    in_body = False
    for block in [b.strip() for b in markup_text.split("\n\n") if b.strip()]:
        if not in_body and block.startswith("# ") and CHAPTER_HEADING_RE.match(block[2:].strip()):
            in_body = True
        if in_body:
            body.append(block)
        else:
            front.append(block)
    return front, body


def _count_front_matter_pages(markup_text: str) -> int:
    front_blocks, _ = _split_front_and_body(markup_text)
    heading_pages = sum(1 for block in front_blocks if block.startswith("# "))
    return heading_pages + 1  # add one TOC page baseline


def _make_paragraph(text: str, style: ParagraphStyle, toc_level: int | None = None) -> Paragraph:
    para = Paragraph(html.escape(text), style)
    if toc_level is not None:
        para.toc_level = toc_level
        para.toc_text = text
    return para


def _build_toc(styles: dict[str, ParagraphStyle]) -> list:
    toc = TableOfContents()
    toc.levelStyles = [styles["toc1"], styles["toc2"], styles["toc3"]]
    toc.dotsMinLevel = 0
    return [
        _make_paragraph("\u76ee\u5f55", styles["toc_title"]),
        Spacer(1, 0.5 * cm),
        toc,
        PageBreak(),
    ]


def _body_story(blocks: list[str], styles: dict[str, ParagraphStyle]) -> list:
    story: list = []
    first_block = True
    main_body_started = False

    for index, block in enumerate(blocks):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        first = lines[0]
        if first.startswith("### "):
            text = first[4:]
            story.append(_make_paragraph(text, styles["h3_cn"], toc_level=2))
        elif first.startswith("## "):
            text = first[3:]
            story.append(_make_paragraph(text, styles["h2_cn"], toc_level=1))
        elif first.startswith("# "):
            heading = first[2:].strip()
            is_chapter = bool(CHAPTER_HEADING_RE.match(heading))
            is_english = _is_mostly_english(heading)
            if is_chapter and not main_body_started:
                story.append(NextPageTemplate("Body"))
                main_body_started = True
            if not first_block:
                story.append(PageBreak())
            style = styles["h1_en"] if is_english else styles["h1_cn"]
            toc_level = 0 if (is_chapter or heading in ["\u7ed3\u8bba", "\u53c2\u8003\u6587\u732e", "\u81f4\u8c22", "\u9644\u5f55"]) else None
            story.append(_make_paragraph(heading, style, toc_level=toc_level))
        else:
            if _is_keyword_line(first):
                style = styles["body_en_no_indent"] if _is_mostly_english(first) else styles["body_cn_no_indent"]
            elif index + 1 < len(blocks) and blocks[index + 1].strip() == "# Abstract" and _looks_like_english_title(block):
                style = styles["center_plain_en"]
            else:
                style = styles["body_en"] if _is_mostly_english(block) and not _contains_cjk(block) else styles["body_cn"]
            text = "<br/>".join(html.escape(line) for line in lines)
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.2 * cm))

        first_block = False

    return story


def build_pdf(markup_text: str, options: FormatOptions, output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    styles = _styles()
    front_pages = _count_front_matter_pages(markup_text)
    cover_pages = 1 if options.include_cover else 0

    doc = ThesisDocTemplate(
        str(output),
        pagesize=A4,
        leftMargin=3.0 * cm,
        rightMargin=2.5 * cm,
        topMargin=3.3 * cm,
        bottomMargin=2.7 * cm,
        title=options.thesis_title or "\u8bba\u6587\u6392\u7248\u9884\u89c8",
        author=options.author_name or "",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    cover_template = PageTemplate(id="Cover", frames=[frame])
    front_template = PageTemplate(id="Front", frames=[frame], onPage=lambda c, d: _front_page(c, d, options.header_text, cover_pages))
    body_template = PageTemplate(id="Body", frames=[frame], onPage=lambda c, d: _body_page(c, d, options.header_text, cover_pages, front_pages))

    if options.include_cover:
        doc.addPageTemplates([cover_template, front_template, body_template])
    else:
        doc.addPageTemplates([front_template, body_template])

    story: list = []
    if options.include_cover:
        story.extend(_cover_story(options, styles))
        story.append(NextPageTemplate("Front"))
        story.append(PageBreak())

    front_blocks, body_blocks = _split_front_and_body(markup_text)
    doc._heading_toc_enabled = False
    story.extend(_body_story(front_blocks, styles))
    if front_blocks:
        story.append(PageBreak())
    story.extend(_build_toc(styles))
    doc._heading_toc_enabled = True
    story.extend(_body_story(body_blocks, styles))
    if not story:
        story.append(Paragraph(" ", styles["body_cn"]))

    doc.multiBuild(story)
    return output


def _set_run_font(run, east_asia: str, ascii_font: str | None = None, size: float | None = None, bold: bool | None = None):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    run.font.name = ascii_font or east_asia
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font or east_asia)
    r_fonts.set(qn("w:hAnsi"), ascii_font or east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_run_color(run, r: int = 0, g: int = 0, b: int = 0) -> None:
    run.font.color.rgb = RGBColor(r, g, b)


def _set_cell_bottom_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def _set_cell_text(cell, text: str, east_font: str, ascii_font: str, size: float, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text or " ")
    _set_run_font(run, east_font, ascii_font, size, bold)
    _set_run_color(run)


def _set_paragraph_border_bottom(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def _clear_story_container(container) -> None:
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def _insert_field(paragraph, instruction: str, display_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if display_text:
        hint = paragraph.add_run(display_text)
        _set_run_font(hint, "SimSun", "Times New Roman", 12, False)
    run = paragraph.add_run()
    run._r.append(end)


def _apply_section_layout(section) -> None:
    section.top_margin = Cm(3.3)
    section.bottom_margin = Cm(2.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def _set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def _set_section_header(section, header_text: str) -> None:
    section.header.is_linked_to_previous = False
    _clear_story_container(section.header)
    p = section.header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(header_text)
    _set_run_font(r, "SimSun", "Times New Roman", 9, False)
    _set_paragraph_border_bottom(p)


def _set_section_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    _clear_story_container(section.footer)
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _insert_field(p, " PAGE ")


def _configure_cover_section(section) -> None:
    _apply_section_layout(section)
    section.different_first_page_header_footer = True
    _clear_story_container(section.header)
    _clear_story_container(section.footer)


def _configure_numbered_section(section, header_text: str, fmt: str, start: int = 1) -> None:
    _apply_section_layout(section)
    section.different_first_page_header_footer = False
    _set_page_number_format(section, fmt, start)
    _set_section_header(section, header_text)
    _set_section_footer(section)


def _configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in (("Heading 1", 15), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.base_style = normal
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(15.6)
        style.paragraph_format.space_after = Pt(15.6)
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_heading(doc: Document, text: str, level: int, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    if _is_mostly_english(text) and not _contains_cjk(text):
        font_name = east = "Times New Roman"
    else:
        font_name = east = "SimHei"
    run = p.add_run(text)
    _set_run_font(run, east, font_name, {1: 15, 2: 14, 3: 12}[level], True)
    _set_run_color(run)


def _add_body_paragraph(doc: Document, text: str, level: int, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    if _is_mostly_english(text) and not _contains_cjk(text):
        font_name = east = "Times New Roman"
    else:
        font_name = east = "SimHei"
    run = p.add_run(text)
    _set_run_font(run, east, font_name, {1: 15, 2: 14, 3: 12}[level], True)


def _add_body_paragraph(doc: Document, lines: list[str], *, no_indent: bool = False, centered: bool = False, hanging: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if hanging:
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0 if no_indent or centered else 0.84)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        if _is_mostly_english(line) and not _contains_cjk(line):
            _set_run_font(run, "Times New Roman", "Times New Roman", 12, False)
        else:
            _set_run_font(run, "SimSun", "Times New Roman", 12, False)
        if idx != len(lines) - 1:
            run.add_break()


def _docx_cover(doc: Document, options: FormatOptions) -> None:
    if options.school_name == "苏州大学" and UNIVERSITY_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(UNIVERSITY_IMAGE), width=Cm(6.8))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(options.school_name or "苏州大学")
        _set_run_font(r, "SimHei", "SimHei", 24, True)
        _set_run_color(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(42)
    title_text = options.header_text.replace(options.school_name, "").strip("（）() ") or "本科毕业设计（论文）"
    r = p.add_run(title_text)
    _set_run_font(r, "SimHei", "SimHei", 22, True)
    _set_run_color(r)

    table = doc.add_table(rows=8, cols=5)
    table.autofit = False
    widths = [Cm(2.8), Cm(3.0), Cm(1.8), Cm(2.6), Cm(4.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    table.cell(0, 1).merge(table.cell(0, 4))
    table.cell(1, 1).merge(table.cell(1, 4))
    table.cell(2, 1).merge(table.cell(2, 2))
    table.cell(3, 1).merge(table.cell(3, 2))
    table.cell(4, 1).merge(table.cell(4, 4))
    table.cell(5, 1).merge(table.cell(5, 2))
    table.cell(6, 1).merge(table.cell(6, 4))

    labels = [
        (0, 0, "学院(部)"),
        (1, 0, "题   目"),
        (2, 0, "年  级"),
        (2, 3, "专业"),
        (3, 0, "班  级"),
        (3, 3, "学号"),
        (4, 0, "姓  名"),
        (5, 0, "指导老师"),
        (5, 3, "职称"),
        (6, 0, "论文提交日期"),
    ]
    values = [
        (0, 1, options.college, WD_ALIGN_PARAGRAPH.CENTER),
        (1, 1, options.thesis_title, WD_ALIGN_PARAGRAPH.CENTER),
        (2, 1, options.year_grade, WD_ALIGN_PARAGRAPH.CENTER),
        (2, 4, options.major, WD_ALIGN_PARAGRAPH.CENTER),
        (3, 1, options.class_name, WD_ALIGN_PARAGRAPH.CENTER),
        (3, 4, options.student_id, WD_ALIGN_PARAGRAPH.CENTER),
        (4, 1, options.author_name, WD_ALIGN_PARAGRAPH.CENTER),
        (5, 1, options.supervisor, WD_ALIGN_PARAGRAPH.LEFT),
        (5, 4, options.supervisor_title, WD_ALIGN_PARAGRAPH.CENTER),
        (6, 1, options.date_text, WD_ALIGN_PARAGRAPH.CENTER),
    ]

    for row_idx, col_idx, label in labels:
        _set_cell_text(table.cell(row_idx, col_idx), label, "SimSun", "SimSun", 14, False, WD_ALIGN_PARAGRAPH.LEFT)

    for row_idx, col_idx, value, align in values:
        _set_cell_text(table.cell(row_idx, col_idx), value or " ", "SimHei", "SimHei", 14, False, align)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = tc_borders.find(qn(f"w:{edge}"))
                if border is None:
                    border = OxmlElement(f"w:{edge}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "nil")

    for row_idx, col_idx, _, _ in values:
        _set_cell_bottom_border(table.cell(row_idx, col_idx))

    doc.add_paragraph("")


def _chapter_number_from_heading(text: str, fallback: int) -> int:
    if CHAPTER_HEADING_RE.match(text):
        return fallback + 1
    match = NUMERIC_CHAPTER_RE.match(text)
    if match:
        return int(match.group(1))
    return fallback + 1


def _renumber_caption(text: str, chapter_index: int, counter: int, kind: str) -> tuple[str, int]:
    counter += 1
    if kind == "figure":
        prefix = f"Figure {chapter_index}-{counter} " if text.lower().startswith("figure") else f"图{chapter_index}-{counter} "
        return FIGURE_CAPTION_RE.sub(prefix, text, count=1), counter
    prefix = f"Table {chapter_index}-{counter} " if text.lower().startswith("table") else f"表{chapter_index}-{counter} "
    return TABLE_CAPTION_RE.sub(prefix, text, count=1), counter


def _renumber_body_blocks(blocks: list[str]) -> list[str]:
    result: list[str] = []
    chapter_index = 0
    figure_count = 0
    table_count = 0
    in_references = False
    reference_index = 0
    body_heading_seen = False

    for block in blocks:
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        first = lines[0]

        if first.startswith("# "):
            heading = first[2:].strip()
            lowered = heading.lower()
            in_references = lowered == "参考文献"
            if lowered not in SPECIAL_HEADINGS and (CHAPTER_HEADING_RE.match(heading) or NUMERIC_CHAPTER_RE.match(heading) or not body_heading_seen):
                chapter_index = _chapter_number_from_heading(heading, chapter_index)
                figure_count = 0
                table_count = 0
                body_heading_seen = True
            result.append(block)
            continue

        if in_references:
            reference_index += 1
            normalized = REFERENCE_PREFIX_RE.sub("", " ".join(lines)).strip()
            result.append(f"[REF]{reference_index} {normalized}")
            continue

        rewritten = list(lines)
        if FIGURE_CAPTION_RE.match(lines[0]):
            if chapter_index == 0:
                chapter_index = 1
            rewritten[0], figure_count = _renumber_caption(lines[0], chapter_index, figure_count, "figure")
        elif TABLE_CAPTION_RE.match(lines[0]):
            if chapter_index == 0:
                chapter_index = 1
            rewritten[0], table_count = _renumber_caption(lines[0], chapter_index, table_count, "table")
        result.append("\n".join(rewritten))

    return result


def _is_table_block(block: str) -> bool:
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    return bool(lines) and all(" | " in line for line in lines)


def _add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            cell = table.cell(row_idx, col_idx)
            cell.text = row[col_idx] if col_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    _set_run_font(run, "SimSun", "Times New Roman", 10.5, False)
    doc.add_paragraph("")


def _insert_formula_placeholder(doc: Document, block: str, formula_map: dict[str, object]) -> bool:
    formula_element = formula_map.get(block.strip())
    if formula_element is None:
        return False
    doc._body._element.append(deepcopy(formula_element))
    return True


def _insert_bound_images(doc: Document, caption_line: str, image_map: dict[str, list]) -> None:
    key = caption_key(caption_line)
    queue = image_map.get(key) or []
    if not queue:
        return
    binding = queue.pop(0)
    width_cm = 12.0
    if binding.width_emu:
        width_cm = max(4.0, min(14.5, binding.width_emu / 360000.0))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(BytesIO(binding.image_bytes), width=Cm(width_cm))


def _build_docx_front_matter(doc: Document, front_blocks: list[str], formula_map: dict[str, object]) -> None:
    index = 0
    while index < len(front_blocks):
        block = front_blocks[index]
        if _insert_formula_placeholder(doc, block, formula_map):
            index += 1
            continue
        if _is_table_block(block):
            rows: list[list[str]] = []
            while index < len(front_blocks) and _is_table_block(front_blocks[index]):
                rows.append([cell.strip() for cell in front_blocks[index].split(" | ")])
                index += 1
            _add_docx_table(doc, rows)
            continue

        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            index += 1
            continue
        first = lines[0]
        if first.startswith("# "):
            _add_heading(doc, first[2:].strip(), 1, page_break_before=index != 0)
        else:
            centered = index + 1 < len(front_blocks) and front_blocks[index + 1].strip() == "# Abstract" and _looks_like_english_title(block)
            _add_body_paragraph(doc, lines, no_indent=_is_keyword_line(first), centered=centered)
        index += 1


def _build_docx_toc(doc: Document) -> None:
    _add_heading(doc, "目录", 1, page_break_before=True)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    _insert_field(p, r'TOC \o "1-3" \h \z \u', '右键单击目录并选择“更新域”以刷新目录。')


def _build_docx_body(doc: Document, body_blocks: list[str], image_map: dict[str, list], formula_map: dict[str, object]) -> None:
    normalized_blocks = _renumber_body_blocks(body_blocks)
    first_body_heading = True
    in_references = False
    index = 0

    while index < len(normalized_blocks):
        block = normalized_blocks[index]
        if block.startswith("[REF]"):
            in_references = True
            marker, ref_body = block.split(" ", 1) if " " in block else (block, "")
            ref_no = marker.replace("[REF]", "")
            _add_body_paragraph(doc, [f"[{ref_no}] {ref_body}".strip()], no_indent=True, hanging=True)
            index += 1
            continue

        if _insert_formula_placeholder(doc, block, formula_map):
            index += 1
            continue

        if _is_table_block(block):
            rows: list[list[str]] = []
            while index < len(normalized_blocks) and _is_table_block(normalized_blocks[index]):
                rows.append([cell.strip() for cell in normalized_blocks[index].split(" | ")])
                index += 1
            _add_docx_table(doc, rows)
            continue

        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            index += 1
            continue
        first = lines[0]

        if first.startswith("# "):
            heading = first[2:].strip()
            in_references = heading.lower() == "参考文献"
            _add_heading(doc, heading, 1, page_break_before=not first_body_heading)
            first_body_heading = False
            index += 1
            continue
        if first.startswith("## "):
            _add_heading(doc, first[3:].strip(), 2)
            index += 1
            continue
        if first.startswith("### "):
            _add_heading(doc, first[4:].strip(), 3)
            index += 1
            continue

        centered = bool(FIGURE_CAPTION_RE.match(first) or TABLE_CAPTION_RE.match(first))
        if centered:
            _insert_bound_images(doc, first, image_map)
        _add_body_paragraph(doc, lines, no_indent=_is_keyword_line(first) or in_references, centered=centered, hanging=in_references)
        index += 1


def build_docx(markup_text: str, options: FormatOptions, output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_docx_styles(doc)
    front_blocks, body_blocks = _split_front_and_body(markup_text)
    image_map = extract_captioned_images(options.source_docx_path) if options.source_docx_path else {}
    formula_map = extract_formula_paragraphs(options.source_docx_path) if options.source_docx_path else {}

    section = doc.sections[0]
    if options.include_cover:
        _configure_cover_section(section)
        _docx_cover(doc, options)
        section = doc.add_section(WD_SECTION.NEW_PAGE)

    _configure_numbered_section(section, options.header_text, "lowerRoman", 1)
    _build_docx_front_matter(doc, front_blocks, formula_map)
    _build_docx_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_numbered_section(body_section, options.header_text, "decimal", 1)
    _build_docx_body(doc, body_blocks, image_map, formula_map)

    doc.save(str(output))
    return output
